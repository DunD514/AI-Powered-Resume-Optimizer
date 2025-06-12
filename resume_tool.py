import pdfplumber
from docx import Document
import re
import nltk
from nltk.tokenize import sent_tokenize, word_tokenize
from nltk.corpus import stopwords
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
import google.generativeai as genai
import os
import time

# Ensure NLTK data
nltk_data_path = os.path.expanduser('~/nltk_data')
if not os.path.exists(nltk_data_path):
    os.makedirs(nltk_data_path)
nltk.data.path.append(nltk_data_path)

try:
    nltk.download('punkt', quiet=True)
    nltk.download('punkt_tab', quiet=True)
    nltk.download('stopwords', quiet=True)
    # Verify tokenization
    test_tokens = word_tokenize("Test Python code.")
    print(f"NLTK tokenization test: {test_tokens}")
except Exception as e:
    print(f"Error downloading NLTK data: {str(e)}")
    print(f"Ensure NLTK data path is writable: {nltk_data_path}")
    print('Run: python -c "import nltk; nltk.download(\'punkt\'); nltk.download(\'punkt_tab\'); nltk.download(\'stopwords\')"')
    print("Or manually download 'punkt_tab' from https://www.nltk.org/nltk_data/ and place in", os.path.join(nltk_data_path, 'tokenizers', 'punkt_tab'))
    exit(1)

# Configure Gemini API
GEMINI_API_KEY = "" # I have removed the API key for security reasons. Please set your own API key.
genai.configure(api_key=GEMINI_API_KEY)

def extract_text_from_pdf(file_path):
    """Extract text from a PDF file."""
    try:
        with pdfplumber.open(file_path) as pdf:
            text = ""
            for page in pdf.pages:
                text += page.extract_text() or ""
            return text
    except Exception as e:
        return f"Error reading PDF: {str(e)}"

def extract_text_from_docx(file_path):
    """Extract text from a DOCX file."""
    try:
        doc = Document(file_path)
        text = "\n".join([para.text for para in doc.paragraphs])
        return text
    except Exception as e:
        return f"Error reading DOCX: {str(e)}"

def extract_key_points(resume_text):
    """Extract key sentences or phrases from resume text."""
    print("Extracted resume text (first 500 chars):\n", resume_text[:500], "...")
    
    try:
        sentences = sent_tokenize(resume_text)
    except Exception as e:
        print(f"Sentence tokenization failed: {str(e)}")
        sentences = resume_text.split('\n')
    
    section_keywords = ['skills', 'technical skills', 'projects', 'certifications', 'education']
    
    key_points = []
    current_section = None
    
    for sentence in sentences:
        sentence = sentence.strip()
        if not sentence or len(sentence) < 10:
            continue
        sentence_lower = sentence.lower()
        for keyword in section_keywords:
            if keyword in sentence_lower and len(sentence.split()) < 15:
                current_section = keyword
                break
        else:
            if current_section in ['skills', 'technical skills', 'projects', 'certifications'] and not any(
                sentence_lower.startswith(x) for x in ['name:', 'email:', 'phone:', 'address:', 'github:', 'linkedin:']
            ):
                key_points.append(sentence)
    
    if len(key_points) < 12:
        print("Insufficient key points. Extracting additional sentences...")
        key_points.extend([s for s in sentences if 10 < len(s.split()) < 50 and not any(
            s.lower().startswith(x) for x in ['name:', 'email:', 'phone:', 'address:', 'github:', 'linkedin:'])][:20-len(key_points)])
    
    print(f"Found {len(key_points)} key points:", key_points)
    return key_points[:20]

def filter_redundant_bullets(bullet_points):
    """Filter out redundant bullet points based on similarity."""
    if len(bullet_points) <= 10:
        return bullet_points
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform(bullet_points)
    similarity_matrix = cosine_similarity(vectors)
    
    keep = []
    used = set()
    for i in range(len(bullet_points)):
        if i not in used:
            keep.append(bullet_points[i])
            for j in range(i + 1, len(bullet_points)):
                if similarity_matrix[i][j] > 0.9:
                    used.add(j)
    return keep[:10]

def generate_bullet_points(resume_text, key_points, job_description):
    """Generate bullet points using Gemini API with retry logic."""
    model = genai.GenerativeModel('gemini-1.5-flash')
    bullet_points = []
    
    for point in key_points:
        prompt = (
            f"Convert this resume sentence into a concise, unique bullet point "
            f"for a Python Software Engineer role at DataNest Technologies. "
            f"Emphasize Python, Flask/FastAPI, Pandas, SQL, AWS, Docker, Pytest, Airflow, GitHub Actions. "
            f"Include specific projects (EchoEstate, Busibot, Velocity X, Detecting Dementia) and quantifiable results (e.g., improved efficiency by 20%). "
            f"Avoid placeholders. Output complete sentences.\n"
            f"Job Description: {job_description[:500]}...\n"
            f"Sentence: {point}\n"
            f"Output a single bullet point starting with a strong action verb, prefixed with '•', max 25 words."
        )
        for attempt in range(4):
            try:
                response = model.generate_content(prompt)
                bullet = response.text.strip()
                if not bullet.startswith(('•', '-', '*')):
                    bullet = f"• {bullet}"
                if '[' not in bullet:
                    bullet_points.append(bullet)
                break
            except Exception as e:
                if "429" in str(e):
                    delay = 45 * (2 ** attempt)
                    print(f"Quota exceeded for point '{point[:50]}...'. Retrying in {delay} seconds...")
                    time.sleep(delay)
                else:
                    print(f"Gemini API error for point '{point[:50]}...': {str(e)}. Using simplified point.")
                    bullet_points.append(f"• {point[:100]}")
                    break
        else:
            bullet_points.append(f"• {point[:100]}")
    
    if len(bullet_points) < 10:
        print(f"Generated only {len(bullet_points)} bullet points. Adding more from resume text...")
        prompt = (
            f"Analyze this resume text and generate {10-len(bullet_points)} concise, unique bullet points "
            f"for a Python Software Engineer role at DataNest Technologies. "
            f"Emphasize Python, Flask/FastAPI, Pandas, SQL, AWS, Docker, Pytest, Airflow, GitHub Actions. "
            f"Include specific projects (EchoEstate, Busibot, Velocity X, Detecting Dementia) and quantifiable results (e.g., improved efficiency by 20%). "
            f"Avoid placeholders. Output complete sentences.\n"
            f"Job Description: {job_description[:500]}...\n"
            f"Resume Text: {resume_text[:2000]}...\n"
            f"Output bullet points starting with strong action verbs, prefixed with '•', max 25 words each."
        )
        for attempt in range(4):
            try:
                response = model.generate_content(prompt)
                generated = response.text.strip().split('\n')
                bullet_points.extend([b.strip() for b in generated if b.strip().startswith('•') and '[' not in b][:(10-len(bullet_points))])
                break
            except Exception as e:
                if "429" in str(e):
                    delay = 45 * (2 ** attempt)
                    print(f"Quota exceeded for resume text. Retrying in {delay} seconds...")
                    time.sleep(delay)
                else:
                    print(f"Gemini API error for resume text: {str(e)}. Using placeholder bullets.")
                    bullet_points.extend([f"• {s[:100]}" for s in resume_text.split('\n') if s.strip() and len(s.split()) < 50][:(10-len(bullet_points))])
                    break
    
    bullet_points = filter_redundant_bullets(bullet_points)
    
    print(f"Generated {len(bullet_points)} bullet points:", bullet_points)
    return bullet_points[:10]

def calculate_keyword_match(resume_text, job_description):
    """Calculate keyword match score."""
    stop_words = set(stopwords.words('english'))
    
    def preprocess(text):
        try:
            tokens = word_tokenize(text.lower())
            tokens = [t for t in tokens if t.isalnum() and t not in stop_words]
            return text
        except Exception as e:
            print(f"Word tokenizeation failed: {str(e)}")
            return " ".join([w.lower() for w in text.split() if w.isalnum()])
    
    resume_clean = preprocess(resume_text)
    job_clean = preprocess(job_description)
    
    vectorizer = TfidfVectorizer()
    vectors = vectorizer.fit_transform([resume_clean, job_clean])
    similarity = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
    return round(similarity * 100, 2)

def validate_file_path(path):
    """Validate file path for output."""
    if not path.endswith('.docx'):
        return False
    try:
        with open(path, 'w') as f:
            pass
        os.remove(path)
        return True
    except Exception:
        return False

def create_new_resume(original_file, bullet_points, output_path):
    """Create a new DOCX resume with generated bullet points."""
    doc = Document(original_file)
    
    target_section = None
    for i, para in enumerate(doc.paragraphs):
        if para.text.strip().lower() == 'projects':
            target_section = i
            break
    
    if target_section is not None:
        doc.paragraphs[target_section].insert_paragraph_before('PROFESSIONAL EXPERIENCE')
        target_section += 1
    else:
        doc.add_paragraph('PROFESSIONAL EXPERIENCE')
        target_section = len(doc.paragraphs)
    
    for bullet in bullet_points:
        doc.paragraphs[target_section].insert_paragraph_before(bullet)
        target_section += 1
    
    doc.save(output_path)
    return output_path

def main():
    try:
        resume_path = input("Enter resume file path (.docx or .pdf): ").strip()
        if not os.path.exists(resume_path):
            raise ValueError("Resume file does not exist.")
        
        choice = input("Do you want to enter a job description file path (y/n)? ").strip().lower()
        if choice not in ['y', 'n']:
            raise ValueError("Invalid input. Please enter 'y' or 'n'.")
        
        if choice == 'y':
            job_desc_path = input("Enter job description file (.docx or .pdf): ").strip()
            if not os.path.exists(job_desc_path):
                raise ValueError("Job description file does not exist.")
            if job_desc_path.endswith('.pdf'):
                job_description = extract_text_from_pdf(job_desc_path)
            elif job_desc_path.endswith('.docx'):
                job_description = extract_text_from_docx(job_desc_path)
            else:
                raise ValueError("Unsupported job description format.")
        else:
            job_description = input("Enter job description text: ").strip()
        
        if resume_path.endswith('.pdf'):
            resume_text = extract_text_from_pdf(resume_path)
        elif resume_path.endswith('.docx'):
            resume_text = extract_text_from_docx(resume_path)
        else:
            raise ValueError("Unsupported resume format.")
        
        key_points = extract_key_points(resume_text)
        bullet_points = generate_bullet_points(resume_text, key_points, job_description)
        
        if not bullet_points:
            raise ValueError("No bullet points generated.")
        
        match_score = calculate_keyword_match(resume_text, job_description)
        print(f"Keyword Match Score: {match_score}%")
        
        while True:
            output_path = input("Enter output file path for rephrased resume (.docx): ").strip()
            if validate_file_path(output_path):
                break
            print("Invalid file path. Please enter a valid .docx file path (e.g., rephrased_resume.docx).")
        
        if resume_path.endswith('.docx'):
            output_file = create_new_resume(resume_path, bullet_points, output_path)
            print(f"Rephrased resume saved to: {output_file}")
        else:
            print("PDF input detected. Generated bullet points:", bullet_points)
    
    except Exception as e:
        print(f"Error: {str(e)}")

if __name__ == "__main__":
    main() 